import io

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from app.schemas.resume_version import ExperienceContent, GeneratedResumeContent


def _experience_heading(item: ExperienceContent) -> str:
    dates = " – ".join(part for part in (item.start_date, item.end_date) if part)
    heading = f"{item.job_title}, {item.employer}"
    return f"{heading} ({dates})" if dates else heading


def generate_docx(content: GeneratedResumeContent) -> bytes:
    """Render assembled resume content as a DOCX document (M7.2).

    Purely a formatting step over already-assembled, human-approved content — no LLM
    involvement and no additional grounding to enforce here.
    """
    document = Document()

    if content.professional_summary:
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(content.professional_summary)

    if content.experiences:
        document.add_heading("Experience", level=1)
        for experience in content.experiences:
            document.add_heading(_experience_heading(experience), level=2)
            if experience.description:
                document.add_paragraph(experience.description)
            for achievement in experience.achievements:
                document.add_paragraph(achievement, style="List Bullet")

    if content.projects:
        document.add_heading("Projects", level=1)
        for project in content.projects:
            document.add_heading(project.name, level=2)
            if project.description:
                document.add_paragraph(project.description)
            if project.technologies:
                document.add_paragraph(f"Technologies: {', '.join(project.technologies)}")
            for achievement in project.achievements:
                document.add_paragraph(achievement, style="List Bullet")

    if content.skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(skill.name for skill in content.skills))

    if content.education:
        document.add_heading("Education", level=1)
        for education in content.education:
            heading = education.institution
            if education.degree:
                heading = f"{education.degree}, {education.institution}"
            document.add_heading(heading, level=2)
            for achievement in education.achievements:
                document.add_paragraph(achievement, style="List Bullet")

    if content.certifications:
        document.add_heading("Certifications", level=1)
        for certification in content.certifications:
            line = certification.name
            if certification.issuing_organization:
                line = f"{line} — {certification.issuing_organization}"
            document.add_paragraph(line, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_pdf(content: GeneratedResumeContent) -> bytes:
    """Render assembled resume content as a PDF document (M7.3), from the same
    already-assembled content used for DOCX export, via reportlab (pure Python, no
    external binary dependency)."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER, title="Resume")
    styles = getSampleStyleSheet()
    heading1 = styles["Heading1"]
    heading2 = styles["Heading2"]
    body = styles["BodyText"]
    bullet_style = ParagraphStyle("Bullet", parent=body, leftIndent=12)

    story: list[object] = []

    if content.professional_summary:
        story.append(Paragraph("Professional Summary", heading1))
        story.append(Paragraph(content.professional_summary, body))
        story.append(Spacer(1, 0.15 * inch))

    if content.experiences:
        story.append(Paragraph("Experience", heading1))
        for experience in content.experiences:
            story.append(Paragraph(_experience_heading(experience), heading2))
            if experience.description:
                story.append(Paragraph(experience.description, body))
            if experience.achievements:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(a, bullet_style)) for a in experience.achievements],
                        bulletType="bullet",
                    )
                )
        story.append(Spacer(1, 0.15 * inch))

    if content.projects:
        story.append(Paragraph("Projects", heading1))
        for project in content.projects:
            story.append(Paragraph(project.name, heading2))
            if project.description:
                story.append(Paragraph(project.description, body))
            if project.technologies:
                story.append(Paragraph(f"Technologies: {', '.join(project.technologies)}", body))
            if project.achievements:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(a, bullet_style)) for a in project.achievements],
                        bulletType="bullet",
                    )
                )
        story.append(Spacer(1, 0.15 * inch))

    if content.skills:
        story.append(Paragraph("Skills", heading1))
        story.append(Paragraph(", ".join(skill.name for skill in content.skills), body))
        story.append(Spacer(1, 0.15 * inch))

    if content.education:
        story.append(Paragraph("Education", heading1))
        for education in content.education:
            heading = education.institution
            if education.degree:
                heading = f"{education.degree}, {education.institution}"
            story.append(Paragraph(heading, heading2))
            if education.achievements:
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(a, bullet_style)) for a in education.achievements],
                        bulletType="bullet",
                    )
                )
        story.append(Spacer(1, 0.15 * inch))

    if content.certifications:
        story.append(Paragraph("Certifications", heading1))
        for certification in content.certifications:
            line = certification.name
            if certification.issuing_organization:
                line = f"{line} — {certification.issuing_organization}"
            story.append(Paragraph(line, body))

    doc.build(story)
    return buffer.getvalue()
